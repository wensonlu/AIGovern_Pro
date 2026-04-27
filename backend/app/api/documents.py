from fastapi import APIRouter, File, UploadFile, Depends, HTTPException, Body
from pydantic import BaseModel
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.models.schemas import DocumentResponse, DocumentListResponse
from app.models.db_models import Document, DocumentChunk
from app.core.config import settings
from app.core.llm import llm_client
from datetime import datetime
import os
import logging
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# 获取日志记录器
logger = logging.getLogger(__name__)


class RetrievalRequest(BaseModel):
    query: str


router = APIRouter(prefix="/api/documents", tags=["documents"])


async def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """从文件提取文本"""
    ext = filename.split(".")[-1].lower()

    try:
        if ext == "txt":
            return file_content.decode("utf-8", errors="replace").replace("\x00", "")
        elif ext == "md":
            return file_content.decode("utf-8", errors="replace").replace("\x00", "")
        elif ext == "pdf":
            # 使用 PyPDF2 提取 PDF 文本
            try:
                pdf_reader = PdfReader(BytesIO(file_content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() or ""
                return text.replace("\x00", "") if text else "[PDF 文本提取失败]"
            except Exception as pdf_err:
                print(f"❌ PDF 提取错误: {pdf_err}")
                return "[PDF 提取失败]"
        elif ext == "docx":
            # 使用 python-docx 提取 DOCX 文本
            try:
                docx_doc = DocxDocument(BytesIO(file_content))
                text = "\n".join([para.text for para in docx_doc.paragraphs])
                return text.replace("\x00", "") if text else "[DOCX 文本提取失败]"
            except Exception as docx_err:
                print(f"❌ DOCX 提取错误: {docx_err}")
                return "[DOCX 提取失败]"
        else:
            return ""
    except Exception as e:
        print(f"❌ 文件提取异常: {e}")
        return "[文件提取异常]"


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """分块处理文本"""
    # 清理 NUL 字符
    text = text.replace("\x00", "").strip()

    if not text:
        return []

    chunks = []
    for i in range(0, len(text), chunk_size - overlap):
        chunk = text[i : i + chunk_size].strip()
        if chunk and len(chunk) > 10:  # 至少10个字符
            chunks.append(chunk)
    return chunks


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    title: str = None,
    category: str = "general",
    db: Session = Depends(get_db),
):
    """上传文档并进行向量化"""

    logger.info(f"[上传接口] 收到上传请求: 文件={file.filename}, 大小={file.size}, 类型={file.content_type}")

    if not title:
        title = file.filename or "Untitled"

    if file.filename:
        ext = file.filename.split(".")[-1].lower()
        logger.info(f"[上传接口] 文件扩展名: {ext}")
        if ext not in settings.allowed_extensions:
            logger.error(f"[上传接口] 不支持的文件类型: {ext}, 允许的类型: {settings.allowed_extensions}")
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {ext}",
            )

    file_content = await file.read()
    logger.info(f"[上传接口] 文件内容读取完成: {len(file_content)} 字节")

    if len(file_content) > settings.max_upload_size:
        logger.error(f"[上传接口] 文件过大: {len(file_content)} > {settings.max_upload_size}")
        raise HTTPException(
            status_code=413,
            detail="文件过大",
        )

    # 保存文件
    logger.info(f"[上传接口] 开始保存文件到 {settings.upload_dir}")
    os.makedirs(settings.upload_dir, exist_ok=True)
    file_path = os.path.join(settings.upload_dir, file.filename or "upload")
    with open(file_path, "wb") as f:
        f.write(file_content)
    logger.info(f"[上传接口] 文件已保存: {file_path}")

    # 提取文本
    logger.info(f"[上传接口] 开始提取文本...")
    text = await extract_text_from_file(file_content, file.filename or "")
    logger.info(f"[上传接口] 文本提取完成: {len(text)} 字符")

    # 分块
    logger.info(f"[上传接口] 开始分块处理...")
    chunks = chunk_text(text)
    logger.info(f"[上传接口] 分块完成: {len(chunks)} 个分块")

    # 创建文档记录
    logger.info(f"[上传接口] 创建文档记录: title={title}, category={category}")
    doc = Document(
        title=title,
        filename=file.filename,  # 保存原始文件名
        category=category,
        embedding_status="processing",
        chunk_count=len(chunks),
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    logger.info(f"[上传接口] 文档记录创建成功: doc_id={doc.id}")

    # 向量化并存入 pgvector
    logger.info(f"[上传接口] 开始向量化 {len(chunks)} 个分块...")
    try:
        for i, chunk in enumerate(chunks):
            embedding = await llm_client.generate_embedding(chunk)

            # 截断到 768 维（Qwen 返回 1536 维，截断到数据库列的维度）
            embedding_768 = embedding[:768] if len(embedding) > 768 else embedding

            # 创建分块记录
            doc_chunk = DocumentChunk(
                document_id=doc.id,
                chunk_index=i,
                chunk_text=chunk,
                embedding=embedding_768,  # pgvector 768维
            )
            db.add(doc_chunk)

            if (i + 1) % 5 == 0:
                logger.debug(f"[上传接口] 向量化进度: {i+1}/{len(chunks)}")

        db.commit()
        embedding_status = "completed"
        logger.info(f"✓ [上传接口] 文档 {doc.id} 向量化完成: {len(chunks)} 个分块 (1536→768维)")
    except Exception as e:
        logger.error(f"⚠️ [上传接口] 向量化失败 (文档 {doc.id}): {e}", exc_info=True)
        embedding_status = "failed"
        db.rollback()

    # 更新文档状态
    doc.embedding_status = embedding_status
    db.commit()

    logger.info(f"[上传接口] 上传完成: doc_id={doc.id}, status={embedding_status}")

    return DocumentResponse(
        id=doc.id,
        title=title,
        category=category,
        embedding_status=embedding_status,
        chunk_count=len(chunks),
        created_at=doc.created_at,
    )


@router.get("", response_model=DocumentListResponse)
async def list_documents(
    skip: int = 0, limit: int = 10, category: str = None, db: Session = Depends(get_db)
):
    """获取文档列表"""
    query = db.query(Document)
    if category:
        query = query.filter(Document.category == category)

    total = query.count()
    items = query.offset(skip).limit(limit).all()

    return DocumentListResponse(
        total=total,
        items=[
            DocumentResponse(
                id=item.id,
                title=item.title,
                category=item.category,
                embedding_status=item.embedding_status,
                chunk_count=item.chunk_count,
                created_at=item.created_at,
            )
            for item in items
        ],
    )


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(document_id: int, db: Session = Depends(get_db)):
    """获取文档详情"""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    return DocumentResponse(
        id=doc.id,
        title=doc.title,
        category=doc.category,
        embedding_status=doc.embedding_status,
        chunk_count=doc.chunk_count,
        created_at=doc.created_at,
    )


@router.delete("/{document_id}")
async def delete_document(document_id: int, db: Session = Depends(get_db)):
    """删除文档"""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 级联删除分块（通过外键配置）
    db.delete(doc)
    db.commit()

    return {"status": "success", "message": "文档已删除"}


@router.post("/{document_id}/test-retrieval")
async def test_retrieval(document_id: int, req: RetrievalRequest, db: Session = Depends(get_db)):
    """测试文档检索效果"""
    from app.services.rag_service import RAGService

    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")

    rag = RAGService(db_session=db)
    results = await rag.retrieve_documents(req.query, top_k=5)

    return {
        "query": req.query,
        "results": results,
        "total": len(results),
    }
